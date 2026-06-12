from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "SA_DEV_TO_PROD_HR_GOVERNANCE_PLAYBOOK.md"
OUTPUT = ROOT / "SA_DEV_TO_PROD_HR_GOVERNANCE_PLAYBOOK.docx"


def set_run_font(run, name="Aptos", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def add_footer(section):
    footer = section.footer
    while len(footer.paragraphs) > 1:
        extra = footer.paragraphs[-1]
        extra._element.getparent().remove(extra._element)
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("PACE ERP | SA Setup Playbook")
    set_run_font(run, size=9, color=(90, 90, 90))


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(18)
    run = p.add_run("PACE ERP")
    set_run_font(run, name="Aptos Display", size=18, bold=True, color=(24, 64, 140))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(10)
    run = p.add_run("SA Dev To Prod-like HR Governance Setup Playbook")
    set_run_font(run, name="Aptos Display", size=22, bold=True, color=(15, 23, 42))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(16)
    run = p.add_run("PDF-ready operator manual for Super Admin")
    set_run_font(run, size=11, color=(71, 85, 105))

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.8)
    rows = [
        ("Prepared For", "Super Admin (SA) operator"),
        ("Document Scope", "Dev-ke Prod-like HR governance structure-e niye jawa"),
        ("Excluded", "Exact user setup, exact company choice, exact approver/viewer person mapping"),
        ("Execution Rule", "Follow sequence exactly. Publish only at the end."),
    ]
    for idx, (left, right) in enumerate(rows):
        table.cell(idx, 0).text = left
        table.cell(idx, 1).text = right
        set_cell_shading(table.cell(idx, 0), "E2E8F0")
        for cell in (table.cell(idx, 0), table.cell(idx, 1)):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=cell == table.cell(idx, 0))

    doc.add_paragraph()
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_heading(doc, text, level):
    if level == 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, name="Aptos Display", size=16, bold=True, color=(20, 83, 45))
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_run_font(run, name="Aptos Display", size=13, bold=True, color=(30, 41, 59))
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        set_run_font(run, size=11.5, bold=True, color=(30, 41, 59))


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.5)


def add_list_paragraph(doc, text, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, size=10.3)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7ED")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.3, bold=True, color=(146, 64, 14))
    doc.add_paragraph()


def parse_markdown_lines(doc, lines):
    number_re = re.compile(r"^\d+\.\s+(.*)$")
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
            continue
        if line.startswith("- "):
            add_list_paragraph(doc, line[2:].strip(), numbered=False)
            continue
        match = number_re.match(line)
        if match:
            add_list_paragraph(doc, match.group(1).strip(), numbered=True)
            continue
        if line.startswith("Important:") or line.startswith("IMPORTANT:"):
            add_callout(doc, line)
            continue
        add_body_paragraph(doc, line)


def build():
    doc = Document()
    configure_page(doc)
    add_cover(doc)
    parse_markdown_lines(doc, SOURCE.read_text(encoding="utf-8").splitlines())
    for section in doc.sections:
        add_footer(section)
    doc.save(OUTPUT)
    print(str(OUTPUT))


if __name__ == "__main__":
    build()
